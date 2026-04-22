#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate วิธี_Push_Code_GitHub.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup (A4) ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── Styles ───────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name='TH Sarabun New', size=None, bold=False, color=None):
    st = styles[style_name]
    st.font.name = font_name
    st.element.rPr.rFonts.set(qn('w:cs'), font_name)
    if size:  st.font.size = Pt(size)
    if bold:  st.font.bold = bold
    if color: st.font.color.rgb = RGBColor(*color)
    return st

# Normal
normal = styles['Normal']
normal.font.name = 'TH Sarabun New'
normal.element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New')
normal.font.size = Pt(14)

# Heading 1
h1 = styles['Heading 1']
h1.font.name = 'TH Sarabun New'
h1.element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New')
h1.font.size = Pt(18)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

# Heading 2
h2 = styles['Heading 2']
h2.font.name = 'TH Sarabun New'
h2.element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New')
h2.font.size = Pt(16)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_para(text='', style='Normal', bold=False, italic=False,
             size=None, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align: p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = 'TH Sarabun New'
        run.font.element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New') if hasattr(run.font.element, 'rPr') else None
        if bold:   run.bold   = bold
        if italic: run.italic = italic
        if size:   run.font.size = Pt(size)
        if color:  run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(lines):
    """Grey-background code block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
        # shade background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

def add_step_box(number, title, commands, description=''):
    """Numbered step with command block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run_num = p.add_run(f'ขั้นตอนที่ {number}  ')
    run_num.font.name = 'TH Sarabun New'
    run_num.font.element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New') if hasattr(run_num.font.element, 'rPr') else None
    run_num.bold = True
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
    run_title = p.add_run(title)
    run_title.font.name = 'TH Sarabun New'
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
    if description:
        add_para(description, size=13, space_before=0, space_after=2)
    add_code_block(commands)

def add_note(text, kind='tip'):
    colors = {'tip': (0xE8, 0xF5, 0xE9), 'warn': (0xFF, 0xF8, 0xE1), 'info': (0xE3, 0xF2, 0xFD)}
    border_colors = {'tip': '2E7D32', 'warn': 'F57F17', 'info': '1565C0'}
    labels = {'tip': '✔ เคล็ดลับ', 'warn': '⚠ ข้อควรระวัง', 'info': 'ℹ หมายเหตุ'}
    label_rgb = {'tip': (0x2E, 0x7D, 0x32), 'warn': (0xF5, 0x7F, 0x17), 'info': (0x15, 0x65, 0xC0)}

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    fill = '%02X%02X%02X' % colors[kind]
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),  'single')
    left.set(qn('w:sz'),   '24')
    left.set(qn('w:space'),'6')
    left.set(qn('w:color'), border_colors[kind])
    pBdr.append(left)
    pPr.append(pBdr)

    run_label = p.add_run(labels[kind] + '  ')
    run_label.font.name = 'TH Sarabun New'
    run_label.bold = True
    run_label.font.size = Pt(13)
    run_label.font.color.rgb = RGBColor(*label_rgb[kind])
    run_text = p.add_run(text)
    run_text.font.name = 'TH Sarabun New'
    run_text.font.size = Pt(13)

def add_hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BDBDBD')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ═══════════════════════════════════════════════════════════════════════════════

add_para('', space_before=0, space_after=20)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(10)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run('คู่มือการ Push Code ขึ้น GitHub')
r.font.name = 'TH Sarabun New'
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
r2 = p_sub.add_run('โปรเจกต์: shortest-path-osm-1')
r2.font.name = 'TH Sarabun New'
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

p_url = doc.add_paragraph()
p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_url.paragraph_format.space_after = Pt(2)
r3 = p_url.add_run('https://github.com/sittinun2tb-2025/shortest-path-osm-1')
r3.font.name = 'Courier New'
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

add_hline()
add_para('', space_before=0, space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: ภาพรวม
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('1. ภาพรวมการทำงาน', level=1)
add_para(
    'Git เป็นระบบ Version Control ที่ติดตามการเปลี่ยนแปลงของไฟล์ '
    'เมื่อแก้ไข code ในเครื่อง (Local) แล้วต้องการส่งขึ้น GitHub (Remote) '
    'จะต้องผ่าน 3 ขั้นตอนหลัก ได้แก่ Stage → Commit → Push',
    size=14, space_after=8
)

# Workflow table
table = doc.add_table(rows=2, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ['Stage (git add)', 'Commit (git commit)', 'Push (git push)']
icons   = ['📁', '💾', '☁']
descs   = [
    'เลือกไฟล์ที่ต้องการบันทึก\n(Staging Area)',
    'บันทึก snapshot ลงใน\nประวัติ Local',
    'ส่ง commit ขึ้น\nGitHub Remote',
]
colors_hex = ['D5E8F0', 'D5F0E0', 'F0E8D5']

for col_idx, (hdr, icon, desc, fill) in enumerate(zip(headers, icons, descs, colors_hex)):
    # Header row
    cell_h = table.rows[0].cells[col_idx]
    cell_h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc = cell_h._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)
    ph = cell_h.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = ph.add_run(f'{icon}  {hdr}')
    rh.font.name = 'TH Sarabun New'; rh.bold = True; rh.font.size = Pt(13)

    # Desc row
    cell_d = table.rows[1].cells[col_idx]
    cell_d.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pd = cell_d.paragraphs[0]
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = pd.add_run(desc)
    rd.font.name = 'TH Sarabun New'; rd.font.size = Pt(12)

add_para('', space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: เตรียมความพร้อม
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('2. การเตรียมความพร้อม', level=1)
add_para('ก่อนเริ่มใช้งาน ตรวจสอบให้แน่ใจว่ามีสิ่งต่อไปนี้:', size=14, space_after=4)

reqs = [
    ('Git', 'ติดตั้งแล้ว — ตรวจสอบด้วย: git --version'),
    ('GitHub Account', 'มีบัญชีและมีสิทธิ์ push ไปยัง repository'),
    ('Repository', 'E:\\Github\\shortest-path-osm-1 เชื่อมต่อกับ remote แล้ว'),
    ('Authentication', 'ตั้งค่า Personal Access Token หรือ SSH Key แล้ว'),
]

req_table = doc.add_table(rows=len(reqs)+1, cols=2)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.LEFT

hrow = req_table.rows[0]
for ci, hd in enumerate(['รายการ', 'รายละเอียด']):
    c = hrow.cells[ci]
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'1565C0')
    tcPr.append(shd)
    rr = c.paragraphs[0].add_run(hd)
    rr.font.name = 'TH Sarabun New'; rr.bold = True; rr.font.size = Pt(13)
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for ri, (item, detail) in enumerate(reqs, 1):
    fill = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
    for ci, text in enumerate([item, detail]):
        c = req_table.rows[ri].cells[ci]
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
        run = c.paragraphs[0].add_run(('✔  ' if ci == 0 else '') + text)
        run.font.name = 'TH Sarabun New'; run.font.size = Pt(13)
        if ci == 0: run.bold = True

add_para('', space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: ขั้นตอน
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('3. ขั้นตอนการ Push Code', level=1)

# Step 1
add_step_box(1, 'เปิด Terminal แล้วไปที่ folder โปรเจกต์',
    ['cd E:/Github/shortest-path-osm-1'],
    'เปิด Command Prompt, PowerShell, หรือ Git Bash แล้วพิมพ์คำสั่งด้านล่าง:')

# Step 2
add_step_box(2, 'ตรวจสอบสถานะไฟล์ที่เปลี่ยนแปลง',
    ['git status'],
    'คำสั่งนี้จะแสดงรายชื่อไฟล์ที่แก้ไข เพิ่ม หรือลบ:')

add_note('ไฟล์สีแดง = ยังไม่ได้ Stage  |  ไฟล์สีเขียว = Stage แล้ว พร้อม Commit', 'info')

# Step 3
add_step_box(3, 'เลือกไฟล์ที่ต้องการ Push (Stage)',
    [
        '# เพิ่มทีละไฟล์',
        'git add params.py',
        'git add run-3-dijkstra.py',
        '',
        '# หรือเพิ่มทุกไฟล์ที่เปลี่ยนแปลงพร้อมกัน',
        'git add .',
    ],
    'เลือกไฟล์ที่ต้องการบันทึก:')

add_note(
    'หลีกเลี่ยง git add . หากมีไฟล์ขนาดใหญ่ เช่น osm_graph.gpkg '
    'ซึ่งถูก ignore ไว้ใน .gitignore แล้ว แต่ควรระวังไฟล์ .pkl ขนาดใหญ่ด้วย',
    'warn'
)

# Step 4
add_step_box(4, 'Commit พร้อมข้อความอธิบาย',
    ['git commit -m "อธิบายสิ่งที่แก้ไข เช่น Update params and dijkstra script"'],
    'บันทึก snapshot ของการเปลี่ยนแปลงลงใน Local repository:')

add_note('ข้อความ commit ควรกระชับ ชัดเจน บอกว่าแก้ไขอะไรและทำไม', 'tip')

# Step 5
add_step_box(5, 'Push ขึ้น GitHub',
    ['git push origin master'],
    'ส่ง commit ทั้งหมดที่อยู่ใน Local ขึ้นไปยัง GitHub:')

add_para('', space_after=8)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: คำสั่งรวม
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('4. คำสั่งรวมทุกขั้นตอน (Quick Reference)', level=1)
add_para('สามารถรันคำสั่งเหล่านี้ต่อเนื่องกันได้เลย:', size=14, space_after=4)

add_code_block([
    'cd E:/Github/shortest-path-osm-1',
    'git status',
    'git add .',
    'git commit -m "อธิบายสิ่งที่แก้ไข"',
    'git push origin master',
])

add_para('', space_after=6)
add_note('หลังจาก push แล้ว สามารถตรวจสอบผลได้ที่\nhttps://github.com/sittinun2tb-2025/shortest-path-osm-1', 'tip')
add_para('', space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: ตรวจสอบผล
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('5. ตรวจสอบผลลัพธ์หลัง Push', level=1)

checks = [
    ('ตรวจสอบ commit ล่าสุดใน Local',  'git log --oneline -5'),
    ('ตรวจสอบว่า branch ตรงกับ Remote', 'git status'),
    ('ดู commit บน GitHub',             'เปิด Browser ไปที่ GitHub repository'),
]
for title, cmd in checks:
    add_para(f'▸  {title}', bold=True, size=14, space_before=4, space_after=2)
    add_code_block([cmd])

add_para('', space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6: ปัญหาที่พบบ่อย
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('6. ปัญหาที่พบบ่อยและวิธีแก้ไข', level=1)

problems = [
    (
        'error: failed to push — rejected (non-fast-forward)',
        'Remote มี commit ที่ Local ไม่มี',
        ['git pull origin master', 'git push origin master'],
    ),
    (
        'fatal: not a git repository',
        'ยังไม่ได้ cd เข้า folder โปรเจกต์',
        ['cd E:/Github/shortest-path-osm-1'],
    ),
    (
        'nothing to commit, working tree clean',
        'ไม่มีไฟล์ที่เปลี่ยนแปลง — ไม่ต้อง push',
        ['# ไม่ต้องทำอะไร ไฟล์อัปเดตแล้ว'],
    ),
    (
        'Authentication failed',
        'Token หมดอายุหรือไม่ถูกต้อง',
        ['# ไปที่ GitHub Settings → Developer settings → Personal access tokens'],

    ),
]

for err, cause, cmds in problems:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    re = p.add_run(f'❌  {err}')
    re.font.name = 'Courier New'; re.font.size = Pt(11)
    re.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    pc = doc.add_paragraph()
    pc.paragraph_format.space_before = Pt(0)
    pc.paragraph_format.space_after  = Pt(2)
    pc.paragraph_format.left_indent  = Cm(0.5)
    rc = pc.add_run(f'สาเหตุ: {cause}')
    rc.font.name = 'TH Sarabun New'; rc.font.size = Pt(13)
    rc.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    add_code_block(cmds)

add_para('', space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7: ไฟล์ที่ควร / ไม่ควร Push
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('7. ไฟล์ที่ควรและไม่ควร Push', level=1)

ft = doc.add_table(rows=1, cols=2)
ft.style = 'Table Grid'
ft.alignment = WD_TABLE_ALIGNMENT.LEFT

for ci, (hd, fill) in enumerate([('✅  ควร Push', '2E7D32'), ('❌  ไม่ควร Push', 'B71C1C')]):
    c = ft.rows[0].cells[ci]
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)
    rr = c.paragraphs[0].add_run(hd)
    rr.font.name = 'TH Sarabun New'; rr.bold = True
    rr.font.size = Pt(14); rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

should     = ['params.py', 'run-*.py', 'rum-1-graph.py', 'flood-1.gpkg', 'README.md', 'CLAUDE.md', 'Proj_1.qgz', 'docs/*.png', 'output.xlsx']
should_not = ['osm_graph.gpkg  (ใหญ่ 8+ MB)', 'cache/  (OSM API cache)', '__pycache__/', '*.pyc', '.claude/  (Claude Code internal)']

max_rows = max(len(should), len(should_not))
for i in range(max_rows):
    row = ft.add_row()
    fill = 'F1F8E9' if i % 2 == 0 else 'FFFFFF'
    fill2 = 'FFF8E1' if i % 2 == 0 else 'FFFFFF'
    for ci, (items, bg) in enumerate([(should, fill), (should_not, fill2)]):
        c = row.cells[ci]
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), bg)
        tcPr.append(shd)
        text = items[i] if i < len(items) else ''
        rr = c.paragraphs[0].add_run(text)
        rr.font.name = 'Courier New'; rr.font.size = Pt(11)

add_para('', space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ═══════════════════════════════════════════════════════════════════════════════

add_hline()
add_para(
    'สร้างโดย Claude Code  |  โปรเจกต์: shortest-path-osm-1  |  github.com/sittinun2tb-2025/shortest-path-osm-1',
    size=11, color=(0x90, 0x90, 0x90), align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=4, space_after=0
)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = r'E:\Github\shortest-path-osm-1\docs\วิธี_Push_Code_GitHub.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
